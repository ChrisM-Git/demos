"""
Run this script to generate the setup guide as a Word document.
Usage: python generate_doc.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────────

def heading1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x0F, 0x4C, 0x81)

def heading2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x1A, 0x7A, 0xC4)

def heading3(text):
    doc.add_heading(text, level=3)

def body(text):
    doc.add_paragraph(text)

def code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xD6, 0x33, 0x00)
    p.paragraph_format.left_indent = Inches(0.4)

def note(text):
    p = doc.add_paragraph()
    run = p.add_run("Note: " + text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

def warning(text):
    p = doc.add_paragraph()
    run = p.add_run("⚠ " + text)
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x44, 0x00)

def spacer():
    doc.add_paragraph("")


# ── Title ────────────────────────────────────────────────────────────────────

title = doc.add_heading("Copilot Studio Red Team Wrapper", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph("Setup Guide — Palo Alto Networks AI Runtime Security (AIRS)")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

spacer()

body(
    "This wrapper bridges the Palo Alto Networks AI Runtime Security (AIRS) red team tool "
    "with Microsoft Copilot Studio via the Direct Line API. Two setup options are provided: "
    "a local Windows 11 setup using ngrok for testing and demos, and a production setup for customers "
    "hosting the wrapper on their own infrastructure."
)

spacer()

# ── Scan Duration Planning ────────────────────────────────────────────────────

heading1("Planning Your Scan Duration")

body(
    "Before starting, consider how many red team categories you plan to run. "
    "Single category scans typically complete in under 30 minutes. "
    "Full multi-category scans can take 2-4 hours or more. "
    "Your setup method should match your expected scan duration."
)

spacer()

table = doc.add_table(rows=1, cols=4)
table.style = "Light Shading Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "Setup Method"
hdr[1].text = "Session Limit"
hdr[2].text = "Best For"
hdr[3].text = "Cost"

rows = [
    ("Windows 11 + ngrok Free", "2 hours", "Single category testing / demos", "Free"),
    ("Windows 11 + ngrok Personal", "Unlimited", "Multi-category scans from laptop", "$10/month"),
    ("Cloud VM (AWS/Azure/GCP)", "Unlimited", "Full customer scans, production use", "~$5-10 one-time"),
]

for method, limit, best, cost in rows:
    row = table.add_row().cells
    row[0].text = method
    row[1].text = limit
    row[2].text = best
    row[3].text = cost

spacer()
warning(
    "ngrok Free tier has a 2-hour session limit. For full multi-category scans (4+ hours), "
    "use ngrok Personal or deploy to a cloud VM. If ngrok expires mid-scan, the AIRS tool "
    "will lose connectivity and the scan will fail."
)
spacer()
note(
    "The Copilot Studio Direct Line token expires after 60 minutes, but this is handled "
    "automatically — the wrapper fetches a fresh token for every request so long scans "
    "are not affected by token expiry."
)

spacer()

# ── Option 1 ─────────────────────────────────────────────────────────────────

heading1("Option 1: Local Windows 11 Setup (using ngrok)")
body(
    "Use this option when you want to run the wrapper on your laptop for testing or demos. "
    "ngrok creates a secure public tunnel to your local machine — no router or firewall changes required."
)

spacer()
heading2("Prerequisites")
for item in [
    "Python 3.9+ (python.org — check 'Add to PATH' during install)",
    "A Copilot Studio bot with Direct Line Speech channel enabled",
    "A free ngrok account (ngrok.com) — upgrade to Personal for scans over 2 hours",
]:
    doc.add_paragraph(item, style="List Bullet")

spacer()
heading2("Step 1 — Install ngrok")
body("Open PowerShell and run:")
code("winget install ngrok.ngrok")
body("Or download the Windows ZIP from ngrok.com/download, extract it, and add the folder to your PATH.")

spacer()
heading2("Step 2 — Authenticate ngrok")
body("Sign up at ngrok.com, copy your auth token from the dashboard, then run:")
code("ngrok config add-authtoken <your-auth-token>")

spacer()
heading2("Step 3 — Set up the wrapper")
body(r"Place main.py, requirements.txt, and .env.example in a folder (e.g. C:\Users\<username>\copilot_wrapper), then:")
code(r"cd C:\Users\<username>\copilot_wrapper")
code("python -m venv .venv")
code(r".venv\Scripts\Activate.ps1")
note("If you get a script execution error, run: Set-ExecutionPolicy -Scope CurrentUser -ExecutionPolicy RemoteSigned")
code("pip install -r requirements.txt")

spacer()
heading2("Step 4 — Configure your Copilot endpoint")
code("copy .env.example .env")
code("notepad .env")
body("Paste your token endpoint into the .env file:")
code("COPILOT_TOKEN_ENDPOINT=https://<your-env>.environment.api.powerplatform.com/powervirtualagents/botsbyschema/<your-bot>/directline/token?api-version=2022-03-01-preview")
body("How to find your token endpoint:")
for item in [
    "Open Copilot Studio",
    "Go to Settings → Channels → Direct Line Speech",
    "Copy the Token Endpoint URL",
]:
    doc.add_paragraph(item, style="List Number")

spacer()
heading2("Step 5 — Start the wrapper")
code("uvicorn main:app --host 0.0.0.0 --port 8000")
warning(
    "Keep this PowerShell window open for the entire scan. If you close it or press Ctrl+C "
    "the wrapper will stop and AIRS will lose connectivity. Disable sleep/screen lock on your laptop during the scan."
)

spacer()
heading2("Step 6 — Expose with ngrok (new PowerShell window)")
body("Open a second PowerShell window and run:")
code("ngrok http 8000")
body("Copy the public URL shown, e.g.: https://abc123.ngrok-free.dev")
warning(
    "Keep this PowerShell window open as well. You need TWO PowerShell windows running simultaneously — "
    "one for uvicorn and one for ngrok. Closing either will stop the scan."
)
note("ngrok Free tier sessions expire after 2 hours. If your scan runs longer, upgrade to ngrok Personal or use Option 2.")

spacer()
heading2("Step 7 — Configure AIRS Red Team Tool")
body("In the AIRS red team tool select cURL Import and enter:")
code('curl -X POST "https://abc123.ngrok-free.dev/" -H "Content-Type: application/json"')
body("Request JSON:")
code('{\n  "input": "{INPUT}"\n}')
body("Response JSON:")
code('{\n  "output": "{RESPONSE}"\n}')
body("Supports Sessions: No")

spacer()

# ── Option 2 ─────────────────────────────────────────────────────────────────

heading1("Option 2: Customer / Production Setup (hosted server)")
body(
    "Use this option for full multi-category scans or when the customer wants to run the wrapper "
    "on their own infrastructure such as a cloud VM (AWS, Azure, GCP) or existing web server. "
    "This provides a stable, always-on endpoint with no session time limits."
)

spacer()
heading2("Prerequisites")
for item in [
    "A Linux server or cloud VM (AWS, Azure, GCP)",
    "Python 3.9+",
    "A public IP address or domain name",
    "Port 8000 open in your firewall or security group",
    "A Copilot Studio bot with Direct Line Speech channel enabled",
]:
    doc.add_paragraph(item, style="List Bullet")

spacer()
heading2("Step 1 — Copy wrapper files to server")
code("scp -r copilot_wrapper/ user@your-server:/opt/copilot_wrapper/")

spacer()
heading2("Step 2 — Install Python and dependencies")
code("sudo apt update && sudo apt install python3 python3-pip python3-venv -y")
code("cd /opt/copilot_wrapper")
code("python3 -m venv venv")
code("source venv/bin/activate")
code("pip install -r requirements.txt")

spacer()
heading2("Step 3 — Configure your Copilot endpoint")
code("cp .env.example .env")
code("nano .env")
body("Set your token endpoint:")
code("COPILOT_TOKEN_ENDPOINT=https://<your-env>.environment.api.powerplatform.com/powervirtualagents/botsbyschema/<your-bot>/directline/token?api-version=2022-03-01-preview")
body("How to find your token endpoint:")
for item in [
    "Open Copilot Studio",
    "Go to Settings → Channels → Direct Line Speech",
    "Copy the Token Endpoint URL",
]:
    doc.add_paragraph(item, style="List Number")

spacer()
heading2("Step 4 — Run the wrapper")
body("For testing:")
code("uvicorn main:app --host 0.0.0.0 --port 8000")
body("For production (runs in background, auto-restarts):")
code("pip install gunicorn")
code("gunicorn main:app -w 4 -k uvicorn.workers.UvicornWorker --bind 0.0.0.0:8000 --daemon")

spacer()
heading2("Step 5 — Open firewall port")
for item in [
    "AWS: Add inbound rule for port 8000 in your Security Group",
    "Azure: Add inbound port rule in Network Security Group",
    "GCP: Add firewall rule allowing TCP port 8000",
]:
    doc.add_paragraph(item, style="List Bullet")

spacer()
heading2("Step 6 — (Optional) Add HTTPS with nginx")
body("For production use, add SSL via nginx:")
code("sudo apt install nginx certbot python3-certbot-nginx -y")
body("Configure nginx to proxy https://yourdomain.com → http://localhost:8000")

spacer()
heading2("Step 7 — Configure AIRS Red Team Tool")
body("In the AIRS red team tool select cURL Import and enter:")
code('curl -X POST "http://your-server-ip:8000/" -H "Content-Type: application/json"')
body("Or with a domain:")
code('curl -X POST "https://yourdomain.com/" -H "Content-Type: application/json"')
body("Request JSON:")
code('{\n  "input": "{INPUT}"\n}')
body("Response JSON:")
code('{\n  "output": "{RESPONSE}"\n}')
body("Supports Sessions: No")

spacer()

# ── Troubleshooting ───────────────────────────────────────────────────────────

heading1("Troubleshooting")

table = doc.add_table(rows=1, cols=3)
table.style = "Light Shading Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "Error"
hdr[1].text = "Cause"
hdr[2].text = "Fix"

rows = [
    ("COPILOT_TOKEN_ENDPOINT is not set", "Missing .env file", "Copy .env.example to .env and fill in the URL"),
    ("502 Token fetch failed", "Wrong token endpoint URL", "Double-check URL from Copilot Studio Direct Line settings"),
    ("502 Conversation start failed", "Direct Line API issue", "Verify the bot is published in Copilot Studio"),
    ("504 Bot did not respond", "Bot too slow or offline", "Check bot health in Copilot Studio or increase poll timeout"),
    ("Connection refused", "Wrapper not running", "Ensure uvicorn is running on port 8000"),
    ("ngrok URL not working", "ngrok session expired (2hr limit)", "Restart ngrok and update URL in AIRS, or upgrade to ngrok Personal"),
    ("Scan fails after 2 hours", "ngrok Free tier session limit", "Upgrade to ngrok Personal or use a cloud VM (Option 2)"),
    ("Occasional timeout errors every few minutes", "Copilot bot slow on certain prompts", "Normal behavior — slow responses are red team findings"),
]

for error, cause, fix in rows:
    row = table.add_row().cells
    row[0].text = error
    row[1].text = cause
    row[2].text = fix

spacer()

# ── Architecture ──────────────────────────────────────────────────────────────

heading1("Architecture Overview")
code(
    "AIRS Red Team Tool (cloud)\n"
    "        |\n"
    "        | HTTPS POST {\"input\": \"...\"}\n"
    "        v\n"
    "   ngrok / public IP\n"
    "        |\n"
    "        v\n"
    "  FastAPI Wrapper (main.py)\n"
    "        |\n"
    "        |-- GET  Copilot Token Endpoint  --> Microsoft Power Platform\n"
    "        |-- POST Direct Line /conversations --> Bot Framework\n"
    "        |-- POST Direct Line /activities  --> Copilot Studio Bot\n"
    "        |-- GET  Direct Line /activities  --> Poll for bot reply\n"
    "        |\n"
    "        | {\"output\": \"...\"}\n"
    "        v\n"
    "AIRS Red Team Tool (receives response & generates report)"
)

# ── Save ──────────────────────────────────────────────────────────────────────

output = "Copilot_Studio_RedTeam_Setup_Guide.docx"
doc.save(output)
print(f"Document saved: {output}")
